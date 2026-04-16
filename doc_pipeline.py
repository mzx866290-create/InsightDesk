"""
文档处理与向量化管道
支持 PDF、Word、Markdown、CSV 等格式的文档加载、分块、向量化和检索
"""

import logging
import math
import os
import re
import shutil
import sys
import tempfile
import time
from pathlib import Path
from typing import Any, Callable, List, Optional
import pandas as pd
from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    UnstructuredMarkdownLoader,
    CSVLoader,
    TextLoader,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document
from sentence_transformers import CrossEncoder
from dotenv import load_dotenv

load_dotenv()

logger = logging.getLogger(__name__)

KEYWORD_TOKEN_PATTERN = re.compile(r"[A-Za-z0-9_./:-]+|[\u4e00-\u9fff]+")

# 设置标准输出为 UTF-8，避免 Windows 控制台 GBK 编码问题
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass


class DocPipeline:
    _embedding_cache = {}
    _reranker_cache = {}
    """文档处理管道类"""

    def __init__(
        self,
        embedding_model: Optional[str] = None,
        device: Optional[str] = None,
        vector_store_path: Optional[str] = None,
    ):
        """
        初始化文档处理管道

        Args:
            embedding_model: Embedding 模型名称
            device: 计算设备 (cpu/cuda)
            vector_store_path: 向量库持久化路径
        """
        self.embedding_model = embedding_model or os.getenv(
            "EMBEDDING_MODEL", "BAAI/bge-base-zh-v1.5"
        )
        self.device = self._resolve_device(device)
        self.vector_store_path = vector_store_path or os.getenv(
            "VECTOR_STORE_PATH", "./vector_store"
        )

        # 延迟加载: 首次使用时才加载模型,避免启动卡顿
        self._embeddings = None
        self._reranker = None  # 延迟加载 Reranker 模型
        self._reranker_device = self.device

        # 普通文档用 800 字符切分，平衡信息密度和检索精准度
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=100,
            separators=["\n\n", "\n", "。", "；", "!", "?", " ", ""],
            length_function=len,
        )
        # 简历/结构化文档常见章节标题关键词
        self._resume_section_keywords = (
            "工作经历", "工作经验", "项目经历", "项目经验",
            "项目介绍", "项目职责", "项目结果",
            "教育背景", "教育经历", "教育信息", "学历",
            "技能", "专业技能", "核心技能", "技术能力",
            "技能证书",
            "自我评价", "个人介绍", "个人简介", "求职意向",
            "荣誉", "证书", "资质", "培训经历",
            "实习经历", "社会实践", "兴趣爱好",
        )
        self._resume_detect_keywords = (
            "工作经历", "工作经验", "项目经历", "求职意向",
            "教育背景", "应聘", "毕业院校", "专业技能",
            "项目介绍", "项目职责", "项目结果", "技能证书", "自我评价",
            "履历", "简历",
        )

        self.vectorstore: Optional[FAISS] = None

    @staticmethod
    def _path_has_non_ascii(path: str) -> bool:
        try:
            path.encode("ascii")
            return False
        except UnicodeEncodeError:
            return True

    def _should_use_faiss_staging_dir(self) -> bool:
        resolved = str(Path(self.vector_store_path).resolve())
        return sys.platform == "win32" and self._path_has_non_ascii(resolved)

    def _make_faiss_staging_dir(self, prefix: str) -> Path:
        base_dir = Path(tempfile.gettempdir()) / "ai_kb_faiss"
        base_dir.mkdir(parents=True, exist_ok=True)
        return Path(tempfile.mkdtemp(prefix=f"{prefix}_", dir=str(base_dir)))

    def _save_vectorstore_local(self) -> None:
        target_dir = Path(self.vector_store_path)
        target_dir.mkdir(parents=True, exist_ok=True)
        if not self._should_use_faiss_staging_dir():
            self.vectorstore.save_local(str(target_dir))
            return

        staging_dir = self._make_faiss_staging_dir("save")
        logger.info(
            "FAISS save workaround enabled for Windows non-ASCII path: %s",
            target_dir,
        )
        try:
            self.vectorstore.save_local(str(staging_dir))
            for file_name in ("index.faiss", "index.pkl"):
                shutil.copy2(staging_dir / file_name, target_dir / file_name)
        finally:
            shutil.rmtree(staging_dir, ignore_errors=True)

    def _load_vectorstore_local(self) -> FAISS:
        target_dir = Path(self.vector_store_path)
        if not self._should_use_faiss_staging_dir():
            return FAISS.load_local(
                str(target_dir),
                self.embeddings,
                allow_dangerous_deserialization=True,
            )

        staging_dir = self._make_faiss_staging_dir("load")
        logger.info(
            "FAISS load workaround enabled for Windows non-ASCII path: %s",
            target_dir,
        )
        try:
            for file_name in ("index.faiss", "index.pkl"):
                shutil.copy2(target_dir / file_name, staging_dir / file_name)
            return FAISS.load_local(
                str(staging_dir),
                self.embeddings,
                allow_dangerous_deserialization=True,
            )
        finally:
            shutil.rmtree(staging_dir, ignore_errors=True)

    def _resolve_device(self, device: Optional[str]) -> str:
        """优先使用显卡；若不可用则回退到 CPU。"""
        configured = device or os.getenv("EMBEDDING_DEVICE")
        if configured:
            normalized = configured.strip().lower()
            if normalized.startswith("cuda"):
                try:
                    import torch

                    if torch.cuda.is_available():
                        return configured
                    logger.warning(
                        "EMBEDDING_DEVICE=%s，但当前 PyTorch/CUDA 不可用，自动回退到 CPU",
                        configured,
                    )
                    return "cpu"
                except Exception:
                    logger.warning(
                        "检测 CUDA 可用性失败，EMBEDDING_DEVICE=%s 自动回退到 CPU",
                        configured,
                        exc_info=True,
                    )
                    return "cpu"
            return configured

        try:
            import torch

            if torch.cuda.is_available():
                return "cuda"
        except Exception:
            logger.debug("torch/cuda 检测失败，回退到 CPU", exc_info=True)

        return "cpu"

    @property
    def embeddings(self):
        """延迟加载 Embedding 模型"""
        if self._embeddings is None:
            cache_key = (self.embedding_model, self.device)
            cached = self._embedding_cache.get(cache_key)
            if cached is not None:
                self._embeddings = cached
                return self._embeddings

            logger.info(
                "加载 Embedding 模型: %s (设备: %s)", self.embedding_model, self.device
            )
            try:
                self._embeddings = HuggingFaceEmbeddings(
                    model_name=self.embedding_model,
                    model_kwargs={"device": self.device},
                    encode_kwargs={"normalize_embeddings": True},
                )
            except Exception as exc:
                lower = str(exc).lower()
                can_retry_on_cpu = self.device != "cpu" and any(
                    token in lower
                    for token in (
                        "cuda",
                        "out of memory",
                        "torch not compiled with cuda enabled",
                    )
                )
                if not can_retry_on_cpu:
                    raise

                logger.warning(
                    "Embedding 模型在设备 %s 上加载失败，自动回退到 CPU: %s",
                    self.device,
                    exc,
                )
                self.device = "cpu"
                self._embeddings = HuggingFaceEmbeddings(
                    model_name=self.embedding_model,
                    model_kwargs={"device": "cpu"},
                    encode_kwargs={"normalize_embeddings": True},
                )
                cache_key = (self.embedding_model, self.device)
            logger.info("Embedding 模型加载完成")
            self._embedding_cache[cache_key] = self._embeddings
        return self._embeddings

    @property
    def reranker(self):
        """延迟加载 Reranker 模型 (用于二段重排)"""
        if self._reranker is None:
            reranker_model = os.getenv("RERANKER_MODEL", "BAAI/bge-reranker-base")
            self._reranker, self._reranker_device = self._load_reranker(
                reranker_model,
                preferred_device=self._reranker_device or self.device,
            )
        return self._reranker

    def _should_retry_reranker_on_cpu(self, exc: Exception, device: str) -> bool:
        if str(device or "").lower() == "cpu":
            return False
        lower = str(exc).lower()
        return any(
            token in lower
            for token in (
                "cuda",
                "cublas",
                "out of memory",
                "torch not compiled with cuda enabled",
                "device-side assert",
                "not enough memory",
            )
        )

    def _create_reranker(
        self,
        model_name: str,
        device: str,
        *,
        local_files_only: bool,
    ) -> CrossEncoder:
        hf_token = os.getenv("HF_TOKEN") or None
        return CrossEncoder(
            model_name,
            max_length=512,
            device=device,
            local_files_only=local_files_only,
            token=hf_token,
        )

    def _load_reranker(
        self,
        model_name: str,
        preferred_device: str,
    ) -> tuple[CrossEncoder, str]:
        candidate_devices = [str(preferred_device or self.device or "cpu").strip() or "cpu"]
        if candidate_devices[0].lower() != "cpu":
            candidate_devices.append("cpu")

        last_error: Exception | None = None
        for device in candidate_devices:
            cache_key = (model_name, device)
            cached = self._reranker_cache.get(cache_key)
            if cached is not None:
                logger.info("复用已缓存 Reranker: %s (device=%s)", model_name, device)
                return cached, device

            logger.info("加载 Reranker 模型: %s (device=%s)", model_name, device)
            try:
                reranker = self._create_reranker(
                    model_name,
                    device,
                    local_files_only=True,
                )
            except Exception:
                logger.warning(
                    "本地缓存未命中，回退到常规方式加载 Reranker: %s (device=%s)",
                    model_name,
                    device,
                    exc_info=True,
                )
                try:
                    reranker = self._create_reranker(
                        model_name,
                        device,
                        local_files_only=False,
                    )
                except Exception as remote_exc:
                    last_error = remote_exc
                    if self._should_retry_reranker_on_cpu(remote_exc, device):
                        logger.warning(
                            "Reranker 在 %s 上加载失败，自动回退到 CPU: %s",
                            device,
                            remote_exc,
                        )
                        continue
                    raise

            self._reranker_cache[cache_key] = reranker
            logger.info("Reranker 模型加载完成 (device=%s)", device)
            return reranker, device

        if last_error is not None:
            raise last_error
        raise RuntimeError("Failed to load reranker model")

    def _predict_rerank_scores(self, query: str, candidates: List[Document]) -> List[float]:
        pairs = [[query, doc.page_content] for doc in candidates]
        if not pairs:
            return []

        reranker_model = os.getenv("RERANKER_MODEL", "BAAI/bge-reranker-base")
        try:
            raw_scores = self.reranker.predict(pairs)
        except Exception as exc:
            if self._should_retry_reranker_on_cpu(exc, self._reranker_device):
                logger.warning(
                    "Reranker 预测在 %s 上失败，自动回退到 CPU: %s",
                    self._reranker_device,
                    exc,
                )
                self._reranker = None
                self._reranker_device = "cpu"
                self._reranker, self._reranker_device = self._load_reranker(
                    reranker_model,
                    preferred_device="cpu",
                )
                raw_scores = self.reranker.predict(pairs)
            else:
                raise

        return self._normalize_rerank_scores(raw_scores, expected=len(candidates))

    def _normalize_rerank_scores(self, raw_scores: Any, expected: int) -> List[float]:
        if hasattr(raw_scores, "tolist"):
            raw_scores = raw_scores.tolist()
        elif hasattr(raw_scores, "__iter__") and not isinstance(
            raw_scores,
            (str, bytes, list, tuple),
        ):
            raw_scores = list(raw_scores)

        if isinstance(raw_scores, (int, float)):
            raw_scores = [raw_scores]
        elif not isinstance(raw_scores, (list, tuple)):
            raw_scores = []

        normalized: List[float] = []
        for value in raw_scores:
            scalar = value
            if isinstance(value, (list, tuple)):
                scalar = value[0] if value else 0.0
            try:
                score = float(scalar)
            except (TypeError, ValueError):
                score = 0.0
            if not math.isfinite(score):
                score = 0.0
            normalized.append(score)

        if expected > 0 and len(normalized) < expected:
            normalized.extend([0.0] * (expected - len(normalized)))

        return normalized[:expected] if expected > 0 else normalized

    def _is_resume_doc(self, text: str) -> bool:
        """判断文本是否为简历类文档（命中 2 个及以上关键词则认定）"""
        count = sum(1 for kw in self._resume_detect_keywords if kw in text)
        return count >= 2

    def _smart_split(self, docs: List[Document]) -> List[Document]:
        """
        智能分块：简历/结构化文档按语义章节切分，普通文档使用标准切分器（800字）。
        简历中每个章节作为一个独立 chunk，保证检索时能拿到完整的工作/项目/教育信息。
        """
        import re as _re

        result: List[Document] = []
        for doc in docs:
            text = doc.page_content
            source = doc.metadata.get("source", "未知")
            # 判断是否为简历类文档
            if not self._is_resume_doc(text):
                before = len(result)
                result.extend(self.splitter.split_documents([doc]))
                logger.debug(
                    "[SmartSplit] 普通文档 %s → %d 个 chunk（chunk_size=800）",
                    source,
                    len(result) - before,
                )
                continue

            # 简历文档：按章节标题分段
            section_keywords = sorted(
                self._resume_section_keywords, key=len, reverse=True
            )
            kw_pattern = "|".join(_re.escape(kw) for kw in section_keywords)
            heading_section_re = _re.compile(
                r"(?:^|\n)\s*(?:#+\s*|【|■|▶|◆|●|○|·)?(?P<section>"
                + kw_pattern
                + r")[\s：:：\】]*",
                _re.MULTILINE,
            )
            positions = [m.start("section") for m in heading_section_re.finditer(text)]

            # mammoth 可能把标题和正文揉到同一行，此时退化为全文关键词切分
            if len(positions) < 2:
                relaxed_keywords = [kw for kw in section_keywords if len(kw) >= 4]
                relaxed_kw_pattern = "|".join(
                    _re.escape(kw) for kw in relaxed_keywords
                )
                relaxed_section_re = _re.compile(
                    r"(?:^|\n|[ \t]{2,}|[。；;!?！？])\s*(?:#+\s*|【|■|▶|◆|●|○|·)?(?P<section>"
                    + relaxed_kw_pattern
                    + r")[\s：:：\】]*",
                    _re.MULTILINE,
                )
                positions = []
                last_pos = -1
                for match in relaxed_section_re.finditer(text):
                    pos = match.start("section")
                    if pos - last_pos < 6:
                        continue
                    positions.append(pos)
                    last_pos = pos

            if len(positions) < 2:
                # 章节识别不到，降级为标准切分
                before = len(result)
                result.extend(self.splitter.split_documents([doc]))
                logger.info(
                    "[SmartSplit] 简历文档 %s 章节识别不足（仅%d处），降级标准切分 → %d chunk",
                    source,
                    len(positions),
                    len(result) - before,
                )
                continue

            # 第一段：章节开始前的头部信息（姓名、联系方式等）
            head_text = text[: positions[0]].strip()
            if head_text:
                result.append(
                    Document(
                        page_content=head_text,
                        metadata={**doc.metadata, "chunk_type": "resume_header"},
                    )
                )

            # 按章节切分，每个章节内容超过 1600 字时继续用标准切分器细分
            section_count = 0
            for i, pos in enumerate(positions):
                end = positions[i + 1] if i + 1 < len(positions) else len(text)
                section_text = text[pos:end].strip()
                if not section_text:
                    continue
                if len(section_text) <= 1600:
                    result.append(
                        Document(
                            page_content=section_text,
                            metadata={**doc.metadata, "chunk_type": "resume_section"},
                        )
                    )
                    section_count += 1
                else:
                    # 超长章节继续切分，chunk_size 较大以保留上下文
                    sub_splitter = RecursiveCharacterTextSplitter(
                        chunk_size=1400,
                        chunk_overlap=200,
                        separators=["\n\n", "\n", "。", "；", "!", "?", " ", ""],
                        length_function=len,
                    )
                    sub_docs = sub_splitter.split_documents(
                        [Document(page_content=section_text, metadata={**doc.metadata, "chunk_type": "resume_section"})]
                    )
                    result.extend(sub_docs)
                    section_count += len(sub_docs)

            logger.info(
                "[SmartSplit] 简历文档 %s → 头部1块 + %d 个章节块，共 %d chunk",
                source,
                section_count,
                (1 if head_text else 0) + section_count,
            )

        return result

    def _load_xlsx(self, file_path: str) -> List[Document]:
        """
        使用 pandas 解析 Excel 文件（支持多 sheet），每个 sheet 生成若干 Document

        Args:
            file_path: Excel 文件路径

        Returns:
            文档列表
        """
        file_name = Path(file_path).name
        if Path(file_path).suffix.lower() == ".xls":
            raise ValueError(
                f"暂不支持 .xls 文件: {file_name}。请先另存为 .xlsx 后再上传。"
            )
        docs: List[Document] = []
        try:
            xls = pd.ExcelFile(file_path, engine="openpyxl")
        except Exception as e:
            raise ValueError(f"Excel 文件无法打开: {file_name} — {e}") from e

        for sheet_name in xls.sheet_names:
            try:
                df = xls.parse(sheet_name)
            except Exception as e:
                logger.warning("  跳过 sheet '%s'（解析失败: %s）", sheet_name, e)
                continue

            if df.empty:
                continue

            # 每行转为一段文本，列名作为 key
            rows_text = []
            for idx, row in df.iterrows():
                cells = " | ".join(
                    f"{col}: {val}"
                    for col, val in row.items()
                    if pd.notna(val) and str(val).strip()
                )
                if cells:
                    rows_text.append(f"行{idx + 1}: {cells}")

            if not rows_text:
                continue

            content = f"【Sheet: {sheet_name}】\n" + "\n".join(rows_text)
            docs.append(
                Document(
                    page_content=content,
                    metadata={
                        "source": file_name,
                        "file_path": file_path,
                        "sheet_name": sheet_name,
                    },
                )
            )

        if not docs:
            raise ValueError(f"Excel 文件中未读取到任何有效数据: {file_name}")

        return docs

    def _load_docx_with_python_docx(self, file_path: str) -> List[Document]:
        """
        使用 python-docx 解析 Word 文档，完整提取段落、表格、文本框和页眉内容。
        能正确处理简历等复杂排版文档（文本框、页眉信息不再遗漏）。

        Args:
            file_path: docx 文件路径

        Returns:
            文档列表（整篇作为一个 Document）
        """
        import docx  # python-docx
        from docx.oxml.ns import qn

        file_name = Path(file_path).name
        doc = docx.Document(file_path)
        parts: List[str] = []

        # ── 步骤 1：提取页眉内容（简历头部信息常放在页眉里）──
        seen_header_texts: set = set()
        for section in doc.sections:
            try:
                header = section.header
                if header is None:
                    continue
                for para in header.paragraphs:
                    text = para.text.strip()
                    if text and text not in seen_header_texts:
                        seen_header_texts.add(text)
                        parts.append(text)
                # 页眉里的表格
                for tbl in header.tables:
                    rows_text: List[str] = []
                    for row in tbl.rows:
                        cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if cells_text:
                            rows_text.append(" | ".join(cells_text))
                    if rows_text:
                        combined = "\n".join(rows_text)
                        if combined not in seen_header_texts:
                            seen_header_texts.add(combined)
                            parts.append(combined)
            except Exception:
                pass  # 某些文档没有页眉，跳过

        # ── 步骤 2：收集所有浮动文本框的内容（txbxContent）──
        txbx_node_ids: set = set()
        txbx_t_ids: set = set()  # 记录文本框里 w:t 的 id，避免步骤3重复提取
        for txbx_content in doc.element.body.findall(".//" + qn("w:txbxContent")):
            txbx_node_ids.add(id(txbx_content))
            texts_in_txbx: List[str] = []
            for para in txbx_content.findall(".//" + qn("w:p")):
                para_text = "".join(
                    (node.text or "")
                    for node in para.iter()
                    if node.tag == qn("w:t")
                ).strip()
                if para_text:
                    texts_in_txbx.append(para_text)
                # 记录文本框内所有 w:t 的 id
                for t_node in para.iter():
                    if t_node.tag == qn("w:t"):
                        txbx_t_ids.add(id(t_node))
            if texts_in_txbx:
                parts.append("\n".join(texts_in_txbx))

        # ── 步骤 3：按顺序遍历 body 提取普通段落和表格（排除文本框内已提取的节点）──
        for block in doc.element.body:
            tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

            if tag == "p":
                # 只提取不在文本框里的 w:t
                para_text = "".join(
                    (node.text or "")
                    for node in block.iter()
                    if node.tag == qn("w:t") and id(node) not in txbx_t_ids
                ).strip()
                if para_text:
                    parts.append(para_text)

            elif tag == "tbl":
                # 表格：每行转为 "col1 | col2 | col3" 格式
                rows_text_list: List[str] = []
                for row in block.findall(".//" + qn("w:tr")):
                    cells: List[str] = []
                    for cell in row.findall(".//" + qn("w:tc")):
                        cell_text = "".join(
                            (node.text or "")
                            for node in cell.iter()
                            if node.tag == qn("w:t") and id(node) not in txbx_t_ids
                        ).strip()
                        cells.append(cell_text)
                    row_text = " | ".join(c for c in cells if c)
                    if row_text:
                        rows_text_list.append(row_text)
                if rows_text_list:
                    parts.append("\n".join(rows_text_list))

        full_text = "\n\n".join(p for p in parts if p.strip()).strip()
        if not full_text:
            raise ValueError(f"python-docx 解析结果为空: {file_name}")

        logger.info(
            "python-docx 解析完成: %s，共 %d 字符", file_name, len(full_text)
        )
        return [
            Document(
                page_content=full_text,
                metadata={"source": file_name, "file_path": file_path},
            )
        ]

    def load_file(self, file_path: str) -> List[Document]:
        """
        根据文件类型自动选择 loader 加载文档

        Args:
            file_path: 文件路径

        Returns:
            加载的文档列表
        """
        file_path = str(Path(file_path).resolve())
        ext = Path(file_path).suffix.lower()
        file_name = Path(file_path).name
        logger.info("加载文件: %s", file_name)

        # XLSX / XLS：使用 pandas 自定义加载
        if ext in (".xlsx", ".xls"):
            docs = self._load_xlsx(file_path)
            return docs

        # PDF：单独捕获异常，提供详细日志
        if ext == ".pdf":
            try:
                loader = PyPDFLoader(file_path)
                docs = loader.load()
            except Exception as e:
                logger.error(
                    "PDF 解析失败: %s — %s（%s）", file_name, type(e).__name__, e
                )
                raise ValueError(f"PDF 解析失败: {file_name} — {e}") from e
            for doc in docs:
                doc.metadata["source"] = file_name
                doc.metadata["file_path"] = file_path
            return docs

        # Markdown：优先 UnstructuredMarkdownLoader，失败降级为 TextLoader
        if ext == ".md":
            try:
                loader = UnstructuredMarkdownLoader(file_path)
                docs = loader.load()
                logger.debug(
                    "Markdown 使用 UnstructuredMarkdownLoader 加载: %s", file_name
                )
            except Exception as e:
                logger.warning(
                    "UnstructuredMarkdownLoader 失败(%s)，降级为 TextLoader: %s",
                    e,
                    file_name,
                )
                loader = TextLoader(file_path, encoding="utf-8")
                docs = loader.load()
            for doc in docs:
                doc.metadata["source"] = file_name
                doc.metadata["file_path"] = file_path
            return docs

        # DOCX / DOC：优先用 python-docx（能提取文本框/页眉/表格，简历不丢内容）
        # 失败降级 mammoth，再失败降级 Docx2txtLoader
        if ext in (".docx", ".doc"):
            # 1. python-docx 三步骤（页眉+文本框+段落/表格）— 对简历等复杂排版最完整
            try:
                docs = self._load_docx_with_python_docx(file_path)
                return docs
            except Exception as e:
                logger.warning("python-docx 解析失败(%s)，降级为 mammoth: %s", e, file_name)

            # 2. mammoth 兜底（纯文本提取，不含文本框/形状）
            try:
                import mammoth
                with open(file_path, "rb") as fh:
                    result = mammoth.extract_raw_text(fh)
                text = (result.value or "").strip()
                if not text:
                    raise ValueError("mammoth 返回空文本")
                logger.info("mammoth 解析完成: %s，共 %d 字符", file_name, len(text))
                return [
                    Document(
                        page_content=text,
                        metadata={"source": file_name, "file_path": file_path},
                    )
                ]
            except Exception as e:
                logger.warning("mammoth 解析失败(%s)，降级为 Docx2txtLoader: %s", e, file_name)

            # 3. Docx2txtLoader 兜底
            loader = Docx2txtLoader(file_path)
            docs = loader.load()
            for doc in docs:
                doc.metadata["source"] = file_name
                doc.metadata["file_path"] = file_path
            return docs

        loaders = {
            ".csv": CSVLoader,
            ".txt": TextLoader,
        }

        loader_cls = loaders.get(ext)
        if not loader_cls:
            raise ValueError(f"不支持的文件类型: {ext}")

        loader = loader_cls(file_path)
        docs = loader.load()

        # 为每个文档添加文件名元数据
        for doc in docs:
            doc.metadata["source"] = file_name
            doc.metadata["file_path"] = file_path

        return docs

    def ingest(
        self,
        file_paths: List[str],
        progress_callback: Optional[Callable[[int], None]] = None,
    ) -> int:
        """
        批量导入文档并构建向量索引

        Args:
            file_paths: 文件路径列表

        Returns:
            导入的文档片段数量
        """
        all_docs = []
        failed_files: list[str] = []
        total_files = len(file_paths)

        def report(progress: int) -> None:
            if progress_callback is not None:
                progress_callback(max(0, min(100, progress)))

        if self.vectorstore is None and os.path.exists(self.vector_store_path):
            self.load_store()

        report(5)

        for index, fp in enumerate(file_paths, start=1):
            try:
                file_start = time.perf_counter()
                docs = self.load_file(fp)
                # 使用智能分块替代原有 splitter.split_documents
                chunks = self._smart_split(docs)
                all_docs.extend(chunks)
                file_name = Path(fp).name
                logger.info(
                    "  ✓ %s: %d 个片段 (耗时 %.2fs)",
                    file_name,
                    len(chunks),
                    time.perf_counter() - file_start,
                )
            except Exception as e:
                file_name = Path(fp).name
                logger.error("  ✗ %s: 加载失败 - %s", file_name, e)
                failed_files.append(f"{file_name}: {e}")
                continue

            if total_files:
                report(10 + int(index / total_files * 35))

        if not all_docs:
            detail = "；".join(failed_files[:3]) if failed_files else "未知原因"
            raise ValueError(f"没有成功加载任何文档。失败原因: {detail}")

        logger.info("开始构建向量索引 (共 %d 个片段)", len(all_docs))
        embed_start = time.perf_counter()
        report(55)

        # 分批向量化，每批 32 个 chunk，避免大文件一次性计算导致长时间阻塞
        BATCH_SIZE = 32
        total_chunks = len(all_docs)
        batches = [all_docs[i : i + BATCH_SIZE] for i in range(0, total_chunks, BATCH_SIZE)]
        logger.info("分批向量化: %d 个片段 / %d 批", total_chunks, len(batches))

        for batch_idx, batch in enumerate(batches):
            if self.vectorstore is None:
                self.vectorstore = FAISS.from_documents(batch, self.embeddings)
            else:
                self.vectorstore.add_documents(batch)
            # 更新进度：55% → 90%
            batch_progress = 55 + int((batch_idx + 1) / len(batches) * 35)
            report(batch_progress)
            logger.debug(
                "  批次 %d/%d 完成 (%d 片段)",
                batch_idx + 1,
                len(batches),
                len(batch),
            )

        logger.info("向量化完成，耗时 %.2fs", time.perf_counter() - embed_start)

        # 持久化
        save_start = time.perf_counter()
        self._save_vectorstore_local()
        logger.info(
            "向量索引已保存到: %s (耗时 %.2fs)",
            self.vector_store_path,
            time.perf_counter() - save_start,
        )
        report(100)

        return len(all_docs)

    def load_store(self) -> bool:
        """
        从磁盘加载已有的向量库

        Returns:
            是否加载成功
        """
        if not os.path.exists(self.vector_store_path):
            logger.warning("向量库不存在: %s", self.vector_store_path)
            return False

        try:
            self.vectorstore = self._load_vectorstore_local()
            logger.info("向量库加载成功: %s", self.vector_store_path)
            return True
        except Exception as e:
            logger.error("向量库加载失败: %s", e)
            return False

    def search(self, query: str, k: int = 4) -> List[Document]:
        """
        检索相关文档片段 (纯余弦相似度)

        Args:
            query: 查询文本
            k: 返回的文档数量

        Returns:
            相关文档列表
        """
        if self.vectorstore is None:
            raise ValueError("向量库未初始化,请先调用 load_store() 或 ingest()")

        return self.vectorstore.similarity_search(query, k=k)

    def search_with_scores(self, query: str, k: int = 4) -> List[tuple[Document, float]]:
        """
        检索相关文档片段并保留原始向量分数

        Args:
            query: 查询文本
            k: 返回的文档数量

        Returns:
            (文档, 分数) 列表；FAISS 分数越小表示越接近
        """
        if self.vectorstore is None:
            raise ValueError("向量库未初始化,请先调用 load_store() 或 ingest()")

        try:
            return list(self.vectorstore.similarity_search_with_score(query, k=k))
        except Exception:
            logger.exception("similarity_search_with_score 失败，回退到 similarity_search")
            return [(doc, 0.0) for doc in self.vectorstore.similarity_search(query, k=k)]

    @staticmethod
    def _normalize_feedback_lookup_text(value: Any) -> str:
        return re.sub(r"\s+", " ", str(value or "").strip()).lower()

    def _load_feedback_summary_map(
        self,
        *,
        source_type: str = "doc",
    ) -> dict[tuple[str, str, str], dict[str, Any]]:
        try:
            from chat_store import aggregate_retrieval_feedback_by_source

            summary = aggregate_retrieval_feedback_by_source(source_type=source_type)
        except Exception:
            logger.exception("加载 retrieval feedback 聚合失败 source_type=%s", source_type)
            return {}

        feedback_map: dict[tuple[str, str, str], dict[str, Any]] = {}
        for item in summary:
            key = (
                self._normalize_feedback_lookup_text(item.get("source_type")),
                self._normalize_feedback_lookup_text(item.get("source_title")),
                self._normalize_feedback_lookup_text(item.get("source_url")),
            )
            feedback_map[key] = item
        return feedback_map

    def _feedback_signal_for_doc(
        self,
        doc: Document,
        feedback_map: dict[tuple[str, str, str], dict[str, Any]],
        *,
        source_type: str = "doc",
    ) -> dict[str, Any]:
        metadata = doc.metadata or {}
        source_title = self._normalize_feedback_lookup_text(
            metadata.get("source") or metadata.get("title") or ""
        )
        source_url = self._normalize_feedback_lookup_text(
            metadata.get("url") or metadata.get("source_url") or ""
        )
        lookup_keys = [
            (self._normalize_feedback_lookup_text(source_type), source_title, source_url),
            (self._normalize_feedback_lookup_text(source_type), source_title, ""),
        ]
        for key in lookup_keys:
            signal = feedback_map.get(key)
            if signal:
                return signal
        return {
            "source_type": source_type,
            "source_title": source_title,
            "source_url": source_url,
            "positive_count": 0,
            "negative_count": 0,
            "net_feedback": 0,
            "total_count": 0,
            "last_updated_at": 0.0,
        }

    @staticmethod
    def _feedback_boost(signal: dict[str, Any]) -> float:
        positive_count = int(signal.get("positive_count") or 0)
        negative_count = int(signal.get("negative_count") or 0)
        net_feedback = int(signal.get("net_feedback") or 0)
        boost = positive_count * 0.08 + net_feedback * 0.06 - negative_count * 0.12
        return round(max(-0.45, min(0.35, boost)), 4)

    def _apply_feedback_metadata(
        self,
        doc: Document,
        *,
        feedback_signal: dict[str, Any],
        score: float | None = None,
        score_key: str = "search_score",
    ) -> Document:
        feedback_boost = self._feedback_boost(feedback_signal)
        metadata = dict(doc.metadata or {})
        metadata.update(
            {
                "feedback_boost": feedback_boost,
                "feedback_net": int(feedback_signal.get("net_feedback") or 0),
                "feedback_positive_count": int(feedback_signal.get("positive_count") or 0),
                "feedback_negative_count": int(feedback_signal.get("negative_count") or 0),
            }
        )
        if score is not None:
            metadata[score_key] = round(float(score), 6)
        return Document(page_content=doc.page_content, metadata=metadata)

    @staticmethod
    def _copy_document_with_metadata(doc: Document, **extra_metadata: Any) -> Document:
        metadata = dict(doc.metadata or {})
        metadata.update(extra_metadata)
        return Document(page_content=doc.page_content, metadata=metadata)

    @staticmethod
    def _normalize_keyword_text(text: str) -> str:
        normalized = str(text or "").strip().lower()
        return re.sub(r"\s+", " ", normalized)

    def _extract_keyword_terms(self, text: str) -> List[str]:
        normalized = self._normalize_keyword_text(text)
        if not normalized:
            return []

        terms: list[str] = []
        seen: set[str] = set()
        for raw_token in KEYWORD_TOKEN_PATTERN.findall(normalized):
            token = raw_token.strip()
            if not token:
                continue
            candidates = [token]
            if re.search(r"[\u4e00-\u9fff]", token) and len(token) >= 4:
                candidates.extend(token[index : index + 2] for index in range(len(token) - 1))
            for candidate in candidates:
                candidate = candidate.strip()
                if not candidate:
                    continue
                if len(candidate) < 2 and not re.search(r"[\u4e00-\u9fff]", candidate):
                    continue
                if candidate in seen:
                    continue
                seen.add(candidate)
                terms.append(candidate)
        return terms

    def _all_index_documents(self) -> List[Document]:
        if self.vectorstore is None:
            raise ValueError("向量库未初始化,请先调用 load_store() 或 ingest()")

        docstore = getattr(self.vectorstore, "docstore", None)
        raw_dict = getattr(docstore, "_dict", None)
        if isinstance(raw_dict, dict):
            return [item for item in raw_dict.values() if isinstance(item, Document)]
        return []

    @staticmethod
    def _document_key(doc: Document) -> str:
        metadata = doc.metadata or {}
        stable_id = (
            metadata.get("chunk_id")
            or metadata.get("id")
            or metadata.get("doc_id")
            or metadata.get("source")
        )
        snippet = re.sub(r"\s+", " ", doc.page_content or "")[:200]
        return f"{stable_id or 'unknown'}::{snippet}"

    def _score_keyword_document(self, query: str, doc: Document) -> dict[str, Any]:
        query_text = self._normalize_keyword_text(query)
        content_text = self._normalize_keyword_text(doc.page_content)
        source_text = self._normalize_keyword_text(doc.metadata.get("source", ""))
        title_text = self._normalize_keyword_text(doc.metadata.get("title", ""))
        source_haystack = " ".join(part for part in (source_text, title_text) if part).strip()
        haystack = " ".join(part for part in (source_haystack, content_text) if part).strip()
        terms = self._extract_keyword_terms(query_text)

        matched_terms: list[str] = []
        term_score = 0.0
        source_term_score = 0.0
        for term in terms:
            hit_count = haystack.count(term)
            if hit_count <= 0:
                continue
            matched_terms.append(term)
            term_weight = min(2.8, 0.5 + len(term) * 0.18)
            term_score += min(hit_count, 3) * term_weight

            source_hits = source_haystack.count(term)
            if source_hits > 0:
                source_term_score += min(source_hits, 2) * (term_weight + 0.5)

        matched_terms = list(dict.fromkeys(matched_terms))
        phrase_score = 4.0 if query_text and query_text in content_text else 0.0
        source_phrase_score = 3.0 if query_text and query_text in source_haystack else 0.0
        coverage = len(matched_terms) / max(1, len(terms))
        score = phrase_score + source_phrase_score + term_score + source_term_score + coverage * 2.0
        return {
            "score": round(score, 4),
            "matched_terms": matched_terms,
            "coverage": round(coverage, 4),
            "score_breakdown": {
                "phrase": round(phrase_score, 4),
                "source_phrase": round(source_phrase_score, 4),
                "term_overlap": round(term_score, 4),
                "source_overlap": round(source_term_score, 4),
                "coverage": round(coverage * 2.0, 4),
            },
        }

    def keyword_search(self, query: str, k: int = 4) -> List[Document]:
        """
        基于关键词匹配的轻量检索，适合精确术语/缩写/编号问题
        """
        if self.vectorstore is None:
            raise ValueError("向量库未初始化,请先调用 load_store() 或 ingest()")

        feedback_map = self._load_feedback_summary_map(source_type="doc")
        scored_docs: list[tuple[Document, dict[str, Any], dict[str, Any]]] = []
        for doc in self._all_index_documents():
            score_meta = self._score_keyword_document(query, doc)
            if score_meta["score"] <= 0:
                continue
            feedback_signal = self._feedback_signal_for_doc(doc, feedback_map, source_type="doc")
            feedback_boost = self._feedback_boost(feedback_signal)
            score_meta["score"] = round(score_meta["score"] + feedback_boost, 4)
            score_meta["score_breakdown"] = {
                **score_meta["score_breakdown"],
                "feedback_boost": feedback_boost,
            }
            scored_docs.append((doc, score_meta, feedback_signal))

        scored_docs.sort(key=lambda item: item[1]["score"], reverse=True)
        top_docs: list[Document] = []
        for rank, (doc, score_meta, feedback_signal) in enumerate(
            scored_docs[: max(1, int(k or 1))], start=1
        ):
            top_docs.append(
                self._apply_feedback_metadata(
                    self._copy_document_with_metadata(
                        doc,
                        search_channel="keyword",
                        search_rank=rank,
                        matched_terms=score_meta["matched_terms"],
                        keyword_coverage=score_meta["coverage"],
                        score_breakdown=score_meta["score_breakdown"],
                    ),
                    feedback_signal=feedback_signal,
                    score=score_meta["score"],
                )
            )
        return top_docs

    def hybrid_search(
        self,
        query: str,
        k: int = 4,
        fetch_k: int = 10,
        use_rerank: bool = False,
    ) -> List[Document]:
        """
        混合检索：向量召回 + 关键词召回，然后做融合排序，可选二段重排
        """
        debug_payload = self.debug_retrieval(
            query,
            search_k=k,
            fetch_k=fetch_k,
            retrieval_mode="hybrid",
            use_rerank=use_rerank,
        )
        results = debug_payload.get("top_results") or []
        docs: list[Document] = []
        for item in results:
            docs.append(
                Document(
                    page_content=str(item.get("full_content") or item.get("snippet") or ""),
                    metadata=dict(item.get("metadata") or {}),
                )
            )
        return docs

    def _format_debug_entry(self, doc: Document, *, rank: int) -> dict[str, Any]:
        metadata = dict(doc.metadata or {})
        snippet = re.sub(r"\s+", " ", doc.page_content or "").strip()
        search_score = metadata.get("search_score")
        if search_score is None:
            search_score = metadata.get("rrf_score")
        matched_terms = metadata.get("matched_terms")
        if not isinstance(matched_terms, list):
            matched_terms = []
        feedback_boost = float(metadata.get("feedback_boost", 0.0) or 0.0)
        feedback_net = int(metadata.get("feedback_net", 0) or 0)
        feedback_positive_count = int(metadata.get("feedback_positive_count", 0) or 0)
        feedback_negative_count = int(metadata.get("feedback_negative_count", 0) or 0)

        return {
            "rank": rank,
            "source": metadata.get("source", "未知"),
            "snippet": snippet[:200],
            "full_content": doc.page_content,
            "score": float(search_score or 0.0),
            "channel": metadata.get("search_channel", "semantic"),
            "matched_terms": matched_terms,
            "feedback_boost": feedback_boost,
            "feedback_net": feedback_net,
            "feedback_positive_count": feedback_positive_count,
            "feedback_negative_count": feedback_negative_count,
            "score_breakdown": metadata.get("score_breakdown") or {},
            "metadata": metadata,
        }

    def _semantic_candidates(self, query: str, fetch_k: int) -> List[Document]:
        pairs = self.search_with_scores(query, k=fetch_k)
        feedback_map = self._load_feedback_summary_map(source_type="doc")
        candidates: list[Document] = []
        for rank, (doc, raw_score) in enumerate(pairs, start=1):
            feedback_signal = self._feedback_signal_for_doc(doc, feedback_map, source_type="doc")
            base_score = round(1.0 / (1.0 + max(float(raw_score), 0.0)), 6)
            adjusted_score = round(base_score + self._feedback_boost(feedback_signal), 6)
            candidates.append(
                self._apply_feedback_metadata(
                    self._copy_document_with_metadata(
                        doc,
                        search_channel="semantic",
                        search_rank=rank,
                        # FAISS 距离越小越好，这里转成可比较的正向分值供调试显示。
                        vector_distance=round(float(raw_score), 6),
                        score_breakdown={
                            "semantic_base": base_score,
                            "feedback_boost": self._feedback_boost(feedback_signal),
                        },
                    ),
                    feedback_signal=feedback_signal,
                    score=adjusted_score,
                )
            )
        candidates.sort(
            key=lambda item: float(item.metadata.get("search_score", 0.0) or 0.0),
            reverse=True,
        )
        return candidates

    def _hybrid_candidates(
        self,
        query: str,
        *,
        fetch_k: int,
    ) -> tuple[List[Document], List[Document], List[Document]]:
        semantic_docs = self._semantic_candidates(query, fetch_k)
        keyword_docs = self.keyword_search(query, k=fetch_k)

        fused_map: dict[str, dict[str, Any]] = {}
        for rank, doc in enumerate(semantic_docs, start=1):
            key = self._document_key(doc)
            fused_map[key] = {
                "doc": doc,
                "semantic_rank": rank,
                "semantic_score": doc.metadata.get("search_score", 0.0),
                "keyword_rank": None,
                "keyword_score": 0.0,
            }

        for rank, doc in enumerate(keyword_docs, start=1):
            key = self._document_key(doc)
            if key not in fused_map:
                fused_map[key] = {
                    "doc": doc,
                    "semantic_rank": None,
                    "semantic_score": 0.0,
                    "keyword_rank": rank,
                    "keyword_score": doc.metadata.get("search_score", 0.0),
                }
            else:
                fused_map[key]["keyword_rank"] = rank
                fused_map[key]["keyword_score"] = doc.metadata.get("search_score", 0.0)
                if (doc.metadata.get("search_score", 0.0) or 0.0) > (
                    fused_map[key]["doc"].metadata.get("search_score", 0.0) or 0.0
                ):
                    fused_map[key]["doc"] = self._copy_document_with_metadata(
                        fused_map[key]["doc"],
                        matched_terms=doc.metadata.get("matched_terms", []),
                        keyword_coverage=doc.metadata.get("keyword_coverage", 0.0),
                    )

        fused_docs: list[Document] = []
        for item in fused_map.values():
            semantic_rank = item["semantic_rank"]
            keyword_rank = item["keyword_rank"]
            rrf_score = 0.0
            if semantic_rank is not None:
                rrf_score += 1.0 / (60 + semantic_rank)
            if keyword_rank is not None:
                rrf_score += 1.0 / (60 + keyword_rank)

            fused_docs.append(
                self._copy_document_with_metadata(
                    item["doc"],
                    search_channel="hybrid",
                    rrf_score=round(rrf_score, 6),
                    search_score=round(
                        rrf_score + float(item["doc"].metadata.get("feedback_boost", 0.0) or 0.0),
                        6,
                    ),
                    semantic_rank=semantic_rank,
                    semantic_score=item["semantic_score"],
                    keyword_rank=keyword_rank,
                    keyword_score=item["keyword_score"],
                    score_breakdown={
                        "semantic_rrf": round(
                            1.0 / (60 + semantic_rank), 6
                        )
                        if semantic_rank is not None
                        else 0.0,
                        "keyword_rrf": round(
                            1.0 / (60 + keyword_rank), 6
                        )
                        if keyword_rank is not None
                        else 0.0,
                        "feedback_boost": float(item["doc"].metadata.get("feedback_boost", 0.0) or 0.0),
                    },
                )
            )

        fused_docs.sort(
            key=lambda doc: float(doc.metadata.get("search_score", 0.0) or 0.0),
            reverse=True,
        )
        return semantic_docs, keyword_docs, fused_docs

    def debug_retrieval(
        self,
        query: str,
        *,
        search_k: int = 5,
        fetch_k: int = 10,
        retrieval_mode: str = "semantic",
        use_rerank: bool = False,
    ) -> dict[str, Any]:
        if self.vectorstore is None:
            raise ValueError("向量库未初始化,请先调用 load_store() 或 ingest()")

        safe_top_k = max(1, int(search_k or 1))
        safe_fetch_k = max(safe_top_k, int(fetch_k or safe_top_k))
        mode = str(retrieval_mode or "semantic").strip().lower()
        if mode not in {"semantic", "keyword", "hybrid"}:
            mode = "semantic"

        query_text = str(query or "").strip()
        rewrite_query = query_text
        query_terms = self._extract_keyword_terms(rewrite_query)

        semantic_candidates: List[Document] = []
        keyword_candidates: List[Document] = []
        fused_candidates: List[Document] = []
        top_docs: List[Document] = []
        search_mode = mode

        if mode == "semantic":
            semantic_candidates = self._semantic_candidates(rewrite_query, safe_fetch_k)
            if use_rerank and semantic_candidates:
                try:
                    rerank_scores = self._predict_rerank_scores(rewrite_query, semantic_candidates)
                    ranked_results = sorted(
                        zip(semantic_candidates, rerank_scores),
                        key=lambda item: item[1],
                        reverse=True,
                    )
                    top_docs = [
                        self._copy_document_with_metadata(
                            doc,
                            search_channel="semantic_rerank",
                            search_score=round(
                                float(score) + float(doc.metadata.get("feedback_boost", 0.0) or 0.0),
                                6,
                            ),
                            rerank_score=round(
                                float(score) + float(doc.metadata.get("feedback_boost", 0.0) or 0.0),
                                6,
                            ),
                            score_breakdown={
                                **(doc.metadata.get("score_breakdown") or {}),
                                "rerank": round(float(score), 6),
                                "feedback_boost": float(doc.metadata.get("feedback_boost", 0.0) or 0.0),
                            },
                        )
                        for doc, score in ranked_results[:safe_top_k]
                    ]
                    search_mode = "semantic_rerank"
                except Exception:
                    logger.exception("semantic debug rerank 失败，回退到语义检索候选")
                    top_docs = semantic_candidates[:safe_top_k]
            else:
                top_docs = semantic_candidates[:safe_top_k]

        elif mode == "keyword":
            keyword_candidates = self.keyword_search(rewrite_query, k=safe_fetch_k)
            top_docs = keyword_candidates[:safe_top_k]
            search_mode = "keyword"

        else:
            semantic_candidates, keyword_candidates, fused_candidates = self._hybrid_candidates(
                rewrite_query,
                fetch_k=safe_fetch_k,
            )
            if use_rerank and fused_candidates:
                try:
                    rerank_scores = self._predict_rerank_scores(rewrite_query, fused_candidates)
                    ranked_results = sorted(
                        zip(fused_candidates, rerank_scores),
                        key=lambda item: item[1],
                        reverse=True,
                    )
                    top_docs = [
                        self._copy_document_with_metadata(
                            doc,
                            search_channel="hybrid_rerank",
                            search_score=round(
                                float(score) + float(doc.metadata.get("feedback_boost", 0.0) or 0.0),
                                6,
                            ),
                            rerank_score=round(
                                float(score) + float(doc.metadata.get("feedback_boost", 0.0) or 0.0),
                                6,
                            ),
                            score_breakdown={
                                **(doc.metadata.get("score_breakdown") or {}),
                                "rerank": round(float(score), 6),
                                "feedback_boost": float(doc.metadata.get("feedback_boost", 0.0) or 0.0),
                            },
                        )
                        for doc, score in ranked_results[:safe_top_k]
                    ]
                    search_mode = "hybrid_rerank"
                except Exception:
                    logger.exception("hybrid debug rerank 失败，回退到融合候选")
                    top_docs = fused_candidates[:safe_top_k]
            else:
                top_docs = fused_candidates[:safe_top_k]
                search_mode = "hybrid"

        matched_terms = sorted(
            {
                term
                for doc in top_docs
                for term in (doc.metadata.get("matched_terms") or [])
                if isinstance(term, str) and term.strip()
            }
        )
        unique_sources = {
            str(doc.metadata.get("source", "") or "").strip()
            for doc in top_docs
            if str(doc.metadata.get("source", "") or "").strip()
        }

        return {
            "results_count": len(top_docs),
            "search_mode": search_mode,
            "retrieval_mode": mode,
            "search_k": safe_top_k,
            "top_k": safe_top_k,
            "fetch_k": safe_fetch_k,
            "rewrite_query": rewrite_query,
            "rewrite_applied": False,
            "query_terms": query_terms,
            "top_results": [
                self._format_debug_entry(doc, rank=index)
                for index, doc in enumerate(top_docs, start=1)
            ],
            "semantic_candidates": [
                self._format_debug_entry(doc, rank=index)
                for index, doc in enumerate(semantic_candidates[:safe_fetch_k], start=1)
            ],
            "keyword_candidates": [
                self._format_debug_entry(doc, rank=index)
                for index, doc in enumerate(keyword_candidates[:safe_fetch_k], start=1)
            ],
            "fused_candidates": [
                self._format_debug_entry(doc, rank=index)
                for index, doc in enumerate(fused_candidates[:safe_fetch_k], start=1)
            ],
            "coverage": {
                "unique_sources": len(unique_sources),
                "source_ratio": round(len(unique_sources) / max(1, len(top_docs)), 4),
                "matched_terms": matched_terms,
                "matched_term_count": len(matched_terms),
            },
        }

    def search_with_rerank(
        self, query: str, k: int = 3, fetch_k: int = 10
    ) -> List[Document]:
        """
        二段重排检索: FAISS 粗排 + BGE-Reranker 精排

        流程:
        1. FAISS 检索 Top fetch_k 个候选文档 (余弦相似度粗排)
        2. 使用 BGE-Reranker 对候选文档重新打分 (语义相关性精排)
        3. 返回得分最高的 Top k 个文档

        Args:
            query: 查询文本
            k: 最终返回的文档数量 (默认 3)
            fetch_k: 初次检索的候选文档数量 (默认 10)

        Returns:
            重排后的文档列表 (按相关性降序)
        """
        if self.vectorstore is None:
            raise ValueError("向量库未初始化,请先调用 load_store() 或 ingest()")

        safe_k = max(1, int(k or 1))
        safe_fetch_k = max(safe_k, int(fetch_k or safe_k))
        if safe_fetch_k != fetch_k or safe_k != k:
            logger.warning(
                "Adjusted rerank parameters: k=%s->%d fetch_k=%s->%d",
                k,
                safe_k,
                fetch_k,
                safe_fetch_k,
            )

        # 第一阶段: FAISS 粗排 (检索更多候选)
        candidates = self.vectorstore.similarity_search(query, k=safe_fetch_k)

        if not candidates:
            return []

        try:
            scores = self._predict_rerank_scores(query, candidates)
        except Exception:
            logger.exception("Reranker 失败，回退到 FAISS similarity_search")
            return candidates[:safe_k]

        feedback_map = self._load_feedback_summary_map(source_type="doc")
        scored_candidates: list[tuple[Document, float, dict[str, Any]]] = []
        for doc, score in zip(candidates, scores):
            feedback_signal = self._feedback_signal_for_doc(doc, feedback_map, source_type="doc")
            adjusted_score = round(float(score) + self._feedback_boost(feedback_signal), 6)
            scored_candidates.append((doc, adjusted_score, feedback_signal))

        # 按得分降序排序
        ranked_results = sorted(
            scored_candidates, key=lambda x: x[1], reverse=True
        )

        # 返回 Top k
        top_docs = [
            self._apply_feedback_metadata(
                self._copy_document_with_metadata(
                    doc,
                    search_channel="semantic_rerank",
                    rerank_score=adjusted_score,
                    score_breakdown={
                        "rerank": adjusted_score,
                        "feedback_boost": self._feedback_boost(feedback_signal),
                    },
                ),
                feedback_signal=feedback_signal,
                score=adjusted_score,
            )
            for doc, adjusted_score, feedback_signal in ranked_results[:safe_k]
        ]

        logger.info(
            "[Rerank] query=%s... fetch_k=%d top_k=%d scores=%s",
            query[:30],
            safe_fetch_k,
            safe_k,
            [f"{s:.4f}" for _, s, _ in ranked_results[:safe_k]],
        )
        for i, (doc, score, _) in enumerate(ranked_results[:safe_k], 1):
            source = doc.metadata.get("source", "未知")
            logger.debug("  %d. %s (得分: %.4f)", i, source, score)

        return top_docs

    def get_stats(self) -> dict:
        """
        获取向量库统计信息

        Returns:
            统计信息字典
        """
        if self.vectorstore is None:
            return {"status": "未初始化", "total_docs": 0}

        return {
            "status": "已加载",
            "total_docs": self.vectorstore.index.ntotal,
            "store_path": self.vector_store_path,
        }

    def delete_store(self) -> bool:
        """
        删除向量库磁盘文件并重置内存状态

        Returns:
            是否删除成功
        """
        import shutil

        self.vectorstore = None
        if not os.path.exists(self.vector_store_path):
            return False
        try:
            shutil.rmtree(self.vector_store_path)
            logger.info("向量库已删除: %s", self.vector_store_path)
            return True
        except Exception as e:
            logger.error("向量库删除失败: %s", e)
            return False


if __name__ == "__main__":
    logging.basicConfig(
        level=logging.INFO, format="%(asctime)s [%(name)s] %(levelname)s %(message)s"
    )
    pipeline = DocPipeline()

    test_files = ["test.pdf", "test.docx", "test.md"]
    existing_files = [f for f in test_files if os.path.exists(f)]

    if existing_files:
        count = pipeline.ingest(existing_files)
        logger.info("导入完成: %d 个文档片段", count)

        results = pipeline.search("测试查询", k=3)
        logger.info("检索结果 (%d 条):", len(results))
        for i, doc in enumerate(results, 1):
            logger.info(
                "%d. %s: %s...",
                i,
                doc.metadata.get("source", "未知"),
                doc.page_content[:100],
            )
    else:
        logger.warning("未找到测试文件,请先创建 test.pdf/test.docx/test.md")
