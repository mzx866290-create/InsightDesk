import io
import json
import re
import time
import uuid
from typing import Any

from pydantic import BaseModel, Field

from backend.chat_store import connect_sqlite
from backend.core.storage_runtime import app_database_path


def _now_timestamp() -> float:
    return time.time()


def _normalize_text(value: Any) -> str:
    return " ".join(str(value or "").strip().split())


def _normalize_optional_text(value: Any) -> str | None:
    normalized = _normalize_text(value)
    return normalized or None


class ArtifactRecord(BaseModel):
    artifact_id: str
    session_id: str
    artifact_type: str
    title: str
    status: str = "ready"
    linked_resource_type: str | None = None
    linked_resource_id: str | None = None
    content: dict[str, Any] = Field(default_factory=dict)
    created_at: float = Field(default_factory=_now_timestamp)
    updated_at: float = Field(default_factory=_now_timestamp)


def artifact_export_formats(artifact: ArtifactRecord) -> list[str]:
    if artifact.artifact_type == "report":
        formats = ["md", "docx", "pptx"]
        if report_has_tables(artifact):
            formats.insert(2, "xlsx")
        return formats
    if artifact.artifact_type == "deck":
        return ["pptx"]
    return []


def _report_markdown(artifact: ArtifactRecord) -> str:
    content = artifact.content if isinstance(artifact.content, dict) else {}
    return str(content.get("markdown") or "")


def report_qa_pairs(artifact: ArtifactRecord) -> list[tuple[str, str]]:
    content = artifact.content if isinstance(artifact.content, dict) else {}
    raw_pairs = content.get("qa_pairs") if isinstance(content.get("qa_pairs"), list) else []
    pairs: list[tuple[str, str]] = []
    for item in raw_pairs:
        if not isinstance(item, dict):
            continue
        question = str(item.get("question") or "").strip()
        answer = str(item.get("answer") or "").strip()
        if question or answer:
            pairs.append((question, answer))
    return pairs


def _markdown_body_lines(markdown: str) -> list[str]:
    lines = str(markdown or "").splitlines()
    if lines and lines[0].strip() == "---":
        for index in range(1, len(lines)):
            if lines[index].strip() == "---":
                return lines[index + 1 :]
    return lines


def _strip_markdown_inline(value: Any) -> str:
    text = str(value or "").strip()
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text.strip()


def _split_markdown_table_row(line: str) -> list[str]:
    stripped = str(line or "").strip().strip("|")
    return [_strip_markdown_inline(cell) for cell in stripped.split("|")]


def _is_markdown_table_separator(line: str) -> bool:
    cells = _split_markdown_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def extract_markdown_tables(markdown: str) -> list[list[list[str]]]:
    lines = _markdown_body_lines(markdown)
    tables: list[list[list[str]]] = []
    index = 0
    while index < len(lines):
        current = lines[index].strip()
        next_line = lines[index + 1].strip() if index + 1 < len(lines) else ""
        if "|" not in current or not _is_markdown_table_separator(next_line):
            index += 1
            continue

        table_rows = [_split_markdown_table_row(current)]
        index += 2
        while index < len(lines) and "|" in lines[index]:
            row = lines[index].strip()
            if not row:
                break
            table_rows.append(_split_markdown_table_row(row))
            index += 1
        if table_rows:
            tables.append(table_rows)
    return tables


def report_has_tables(artifact: ArtifactRecord) -> bool:
    return bool(extract_markdown_tables(_report_markdown(artifact)))


def export_report_to_docx(artifact: ArtifactRecord) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed.") from exc

    document = Document()
    document.add_heading(_normalize_text(artifact.title) or "Report", level=0)

    qa_pairs = report_qa_pairs(artifact)
    if qa_pairs:
        for index, (question, answer) in enumerate(qa_pairs, start=1):
            heading = question or f"Question {index}"
            document.add_heading(heading, level=1)
            _append_markdown_to_docx(document, answer)
    else:
        _append_markdown_to_docx(document, _report_markdown(artifact))

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _append_markdown_to_docx(document: Any, markdown: str) -> None:
    pending_table: list[list[str]] = []

    def flush_table() -> None:
        nonlocal pending_table
        if not pending_table:
            return
        max_cols = max(len(row) for row in pending_table)
        table = document.add_table(rows=len(pending_table), cols=max_cols)
        table.style = "Table Grid"
        for row_index, row in enumerate(pending_table):
            for col_index in range(max_cols):
                value = row[col_index] if col_index < len(row) else ""
                table.cell(row_index, col_index).text = value
        pending_table = []

    lines = _markdown_body_lines(markdown)
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()
        next_line = lines[index + 1].strip() if index + 1 < len(lines) else ""

        if line and "|" in line and _is_markdown_table_separator(next_line):
            flush_table()
            pending_table = [_split_markdown_table_row(line)]
            index += 2
            while index < len(lines) and "|" in lines[index]:
                row_line = lines[index].strip()
                if not row_line:
                    break
                pending_table.append(_split_markdown_table_row(row_line))
                index += 1
            flush_table()
            continue

        flush_table()
        if not line:
            index += 1
            continue
        if line.startswith("#"):
            level = min(3, max(1, len(line) - len(line.lstrip("#"))))
            document.add_heading(_strip_markdown_inline(line.lstrip("#").strip()), level=level)
        elif re.match(r"^[-*]\s+", line):
            document.add_paragraph(_strip_markdown_inline(line[2:]), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", line):
            document.add_paragraph(
                _strip_markdown_inline(re.sub(r"^\d+[.)]\s+", "", line)),
                style="List Number",
            )
        else:
            document.add_paragraph(_strip_markdown_inline(line))
        index += 1
    flush_table()


def export_report_to_xlsx(artifact: ArtifactRecord) -> bytes:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
    except ImportError as exc:
        raise RuntimeError("openpyxl is not installed.") from exc

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Report"
    header_fill = PatternFill("solid", fgColor="1F4E79")
    header_font = Font(color="FFFFFF", bold=True)
    wrap = Alignment(wrap_text=True, vertical="top")

    sheet.append(["Title", artifact.title])
    sheet.append([])
    sheet.append(["Question", "Answer"])
    for cell in sheet[3]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = wrap

    for question, answer in report_qa_pairs(artifact):
        sheet.append([question, answer])

    for row in sheet.iter_rows():
        for cell in row:
            cell.alignment = wrap
    sheet.column_dimensions["A"].width = 36
    sheet.column_dimensions["B"].width = 90
    sheet.freeze_panes = "A4"

    for table_index, rows in enumerate(extract_markdown_tables(_report_markdown(artifact)), start=1):
        table_sheet = workbook.create_sheet(f"Table {table_index}")
        for row in rows:
            table_sheet.append(row)
        if rows:
            for cell in table_sheet[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = wrap
        for row in table_sheet.iter_rows():
            for cell in row:
                cell.alignment = wrap
        for column_cells in table_sheet.columns:
            letter = column_cells[0].column_letter
            max_len = max(len(str(cell.value or "")) for cell in column_cells)
            table_sheet.column_dimensions[letter].width = min(max(max_len + 2, 14), 48)

    buffer = io.BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return buffer.read()


def build_report_artifact(
    *,
    session_id: str,
    title: str,
    markdown: str,
    qa_pairs: list[tuple[str, str]],
    answer_group_id: str = "",
    panel_id: str = "",
    artifact_id: str | None = None,
) -> ArtifactRecord:
    now = _now_timestamp()
    return ArtifactRecord(
        artifact_id=artifact_id or f"artifact_{uuid.uuid4().hex}",
        session_id=_normalize_text(session_id),
        artifact_type="report",
        title=_normalize_text(title) or "AI 对话报告",
        status="ready",
        content={
            "markdown": str(markdown or ""),
            "qa_pairs": [
                {
                    "question": _normalize_text(question),
                    "answer": str(answer or "").strip(),
                }
                for question, answer in qa_pairs
                if _normalize_text(question) or str(answer or "").strip()
            ],
            "answer_group_id": _normalize_optional_text(answer_group_id),
            "panel_id": _normalize_optional_text(panel_id),
        },
        created_at=now,
        updated_at=now,
    )


def _first_research_report_artifact(agent_result: dict[str, Any]) -> dict[str, Any]:
    artifacts = agent_result.get("artifacts")
    if not isinstance(artifacts, list):
        return {}
    for artifact in artifacts:
        if not isinstance(artifact, dict):
            continue
        if artifact.get("type") == "research_report" or artifact.get("version") == "v2":
            return dict(artifact)
    return {}


def build_research_archive_artifact(
    *,
    session_id: str,
    title: str,
    agent_result: dict[str, Any],
    task_id: str = "",
    artifact_id: str | None = None,
) -> ArtifactRecord:
    """Persist a Research V2 agent result as a report-compatible archive."""
    now = _now_timestamp()
    research_report = _first_research_report_artifact(agent_result)
    output = str(agent_result.get("output") or "").strip()
    metadata = (
        dict(agent_result.get("metadata") or {})
        if isinstance(agent_result.get("metadata"), dict)
        else {}
    )
    sources = (
        list(agent_result.get("sources") or [])
        if isinstance(agent_result.get("sources"), list)
        else []
    )
    query = _normalize_text(research_report.get("query") or title)
    artifact_title = _normalize_text(title) or query or "Research Archive"
    content: dict[str, Any] = {
        "markdown": output,
        "qa_pairs": [
            {
                "question": query or artifact_title,
                "answer": output,
            }
        ]
        if output
        else [],
        "answer_group_id": None,
        "panel_id": None,
        "research_archive": True,
        "research_report": research_report,
        "sources": sources,
        "metadata": metadata,
        "task_id": _normalize_optional_text(task_id),
        "claim_evidence_chains": list(research_report.get("claim_evidence_chains") or []),
        "claim_verification_summary": research_report.get("claim_verification_summary") or {},
    }
    return ArtifactRecord(
        artifact_id=artifact_id or f"artifact_{uuid.uuid4().hex}",
        session_id=_normalize_text(session_id),
        artifact_type="report",
        title=artifact_title,
        status="ready",
        linked_resource_type="task" if _normalize_text(task_id) else None,
        linked_resource_id=_normalize_optional_text(task_id),
        content=content,
        created_at=now,
        updated_at=now,
    )


def build_deck_artifact(
    deck: Any,
    *,
    artifact_id: str | None = None,
) -> ArtifactRecord:
    now = _now_timestamp()
    content = _deck_artifact_content(deck)
    return ArtifactRecord(
        artifact_id=artifact_id or f"artifact_{uuid.uuid4().hex}",
        session_id=_normalize_text(
            getattr(getattr(deck, "meta", None), "session_id", "")
        ),
        artifact_type="deck",
        title=_normalize_text(getattr(getattr(deck, "meta", None), "title", ""))
        or "Deck Draft",
        status=_normalize_text(getattr(deck, "status", "")) or "draft",
        linked_resource_type="deck",
        linked_resource_id=_normalize_optional_text(getattr(deck, "deck_id", "")),
        content=content,
        created_at=now,
        updated_at=now,
    )


def sync_deck_artifact(artifact: ArtifactRecord, deck: Any) -> ArtifactRecord:
    artifact.session_id = _normalize_text(
        getattr(getattr(deck, "meta", None), "session_id", "")
    )
    artifact.title = (
        _normalize_text(getattr(getattr(deck, "meta", None), "title", ""))
        or artifact.title
    )
    artifact.status = _normalize_text(getattr(deck, "status", "")) or artifact.status
    artifact.linked_resource_type = "deck"
    artifact.linked_resource_id = _normalize_optional_text(getattr(deck, "deck_id", ""))
    artifact.content = _deck_artifact_content(deck)
    artifact.updated_at = _now_timestamp()
    return artifact


def _model_dump_jsonable(value: Any) -> Any:
    if hasattr(value, "model_dump"):
        return value.model_dump(mode="json")
    return value


def _deck_artifact_content(deck: Any) -> dict[str, Any]:
    generation = getattr(deck, "generation", None)
    return {
        "deck_id": _normalize_text(getattr(deck, "deck_id", "")),
        "theme": _normalize_text(getattr(getattr(deck, "meta", None), "theme", ""))
        or "default",
        "slide_count": int(
            getattr(generation, "actual_slide_count", 0) or 0
        ),
        "answer_group_id": _normalize_optional_text(
            getattr(getattr(deck, "meta", None), "source_answer_group_id", "")
        ),
        "panel_id": _normalize_optional_text(
            getattr(getattr(deck, "meta", None), "source_panel_id", "")
        ),
        "evidence_coverage": _model_dump_jsonable(
            getattr(generation, "evidence_coverage", None)
        ),
        "evidence_review": _model_dump_jsonable(
            getattr(generation, "evidence_review", None)
        ),
        "citation_validation": _model_dump_jsonable(
            getattr(deck, "citation_validation", None)
            or getattr(generation, "citation_validation", None)
        ),
    }


class SQLiteArtifactStore:
    def __init__(self, db_path: str | None = None):
        self.db_path = str(db_path or app_database_path()).strip()
        self._init_db()

    def _init_db(self) -> None:
        with connect_sqlite(self.db_path) as conn:
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS artifacts (
                    artifact_id TEXT PRIMARY KEY,
                    session_id TEXT NOT NULL,
                    artifact_type TEXT NOT NULL,
                    title TEXT NOT NULL,
                    status TEXT NOT NULL,
                    linked_resource_type TEXT,
                    linked_resource_id TEXT,
                    content_json TEXT NOT NULL,
                    created_at REAL NOT NULL,
                    updated_at REAL NOT NULL
                )
                """
            )
            conn.execute(
                """
                CREATE INDEX IF NOT EXISTS idx_artifacts_session
                ON artifacts(session_id, updated_at DESC, created_at DESC)
                """
            )
            conn.execute(
                """
                CREATE INDEX IF NOT EXISTS idx_artifacts_linked_resource
                ON artifacts(linked_resource_type, linked_resource_id, updated_at DESC)
                """
            )
            conn.commit()

    def save(self, artifact: ArtifactRecord) -> ArtifactRecord:
        self._init_db()
        now = _now_timestamp()
        created_at = float(getattr(artifact, "created_at", now) or now)
        artifact.updated_at = now
        payload = json.dumps(artifact.content, ensure_ascii=False)
        with connect_sqlite(self.db_path) as conn:
            conn.execute(
                """
                INSERT INTO artifacts (
                    artifact_id,
                    session_id,
                    artifact_type,
                    title,
                    status,
                    linked_resource_type,
                    linked_resource_id,
                    content_json,
                    created_at,
                    updated_at
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ON CONFLICT(artifact_id) DO UPDATE SET
                    session_id = excluded.session_id,
                    artifact_type = excluded.artifact_type,
                    title = excluded.title,
                    status = excluded.status,
                    linked_resource_type = excluded.linked_resource_type,
                    linked_resource_id = excluded.linked_resource_id,
                    content_json = excluded.content_json,
                    updated_at = excluded.updated_at
                """,
                (
                    artifact.artifact_id,
                    artifact.session_id,
                    artifact.artifact_type,
                    artifact.title,
                    artifact.status,
                    artifact.linked_resource_type,
                    artifact.linked_resource_id,
                    payload,
                    created_at,
                    now,
                ),
            )
            conn.commit()
        artifact.created_at = created_at
        artifact.updated_at = now
        return artifact

    def get(self, artifact_id: str) -> ArtifactRecord:
        self._init_db()
        with connect_sqlite(self.db_path) as conn:
            row = conn.execute(
                """
                SELECT
                    artifact_id,
                    session_id,
                    artifact_type,
                    title,
                    status,
                    linked_resource_type,
                    linked_resource_id,
                    content_json,
                    created_at,
                    updated_at
                FROM artifacts
                WHERE artifact_id = ?
                """,
                (str(artifact_id or "").strip(),),
            ).fetchone()
        if not row:
            raise KeyError(str(artifact_id or "").strip())
        return ArtifactRecord(
            artifact_id=str(row[0] or ""),
            session_id=str(row[1] or ""),
            artifact_type=str(row[2] or ""),
            title=str(row[3] or ""),
            status=str(row[4] or "") or "ready",
            linked_resource_type=_normalize_optional_text(row[5]),
            linked_resource_id=_normalize_optional_text(row[6]),
            content=json.loads(row[7] or "{}"),
            created_at=float(row[8] or 0),
            updated_at=float(row[9] or 0),
        )

    def list_recent(
        self,
        *,
        limit: int = 100,
        artifact_type: str = "",
    ) -> list[ArtifactRecord]:
        self._init_db()
        safe_limit = max(1, min(500, int(limit or 100)))
        query = """
            SELECT artifact_id
            FROM artifacts
        """
        params: list[Any] = []
        normalized_artifact_type = _normalize_text(artifact_type)
        if normalized_artifact_type:
            query += " WHERE artifact_type = ?"
            params.append(normalized_artifact_type)
        query += " ORDER BY updated_at DESC, created_at DESC LIMIT ?"
        params.append(safe_limit)
        with connect_sqlite(self.db_path) as conn:
            rows = conn.execute(query, tuple(params)).fetchall()
        return [self.get(str(row[0] or "")) for row in rows if row and row[0]]

    def list_by_session(
        self,
        session_id: str,
        *,
        artifact_type: str = "",
    ) -> list[ArtifactRecord]:
        normalized_session_id = _normalize_text(session_id)
        if not normalized_session_id:
            return []

        self._init_db()
        query = """
            SELECT artifact_id
            FROM artifacts
            WHERE session_id = ?
        """
        params: list[Any] = [normalized_session_id]
        normalized_artifact_type = _normalize_text(artifact_type)
        if normalized_artifact_type:
            query += " AND artifact_type = ?"
            params.append(normalized_artifact_type)
        query += " ORDER BY updated_at DESC, created_at DESC"
        with connect_sqlite(self.db_path) as conn:
            rows = conn.execute(query, tuple(params)).fetchall()
        return [self.get(str(row[0] or "")) for row in rows if row and row[0]]

    def list_by_linked_resource(
        self,
        resource_type: str,
        resource_id: str,
    ) -> list[ArtifactRecord]:
        normalized_resource_type = _normalize_text(resource_type)
        normalized_resource_id = _normalize_text(resource_id)
        if not normalized_resource_type or not normalized_resource_id:
            return []

        self._init_db()
        with connect_sqlite(self.db_path) as conn:
            rows = conn.execute(
                """
                SELECT artifact_id
                FROM artifacts
                WHERE linked_resource_type = ? AND linked_resource_id = ?
                ORDER BY updated_at DESC, created_at DESC
                """,
                (normalized_resource_type, normalized_resource_id),
            ).fetchall()
        return [self.get(str(row[0] or "")) for row in rows if row and row[0]]

    def delete_by_session(self, session_id: str) -> int:
        normalized_session_id = _normalize_text(session_id)
        if not normalized_session_id:
            return 0

        self._init_db()
        with connect_sqlite(self.db_path) as conn:
            cursor = conn.execute(
                "DELETE FROM artifacts WHERE session_id = ?",
                (normalized_session_id,),
            )
            conn.commit()
            return int(cursor.rowcount or 0)
